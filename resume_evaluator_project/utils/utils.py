from pathlib import Path
from pypdf import PdfReader
from docx import Document
import json
from pydantic import ValidationError

def clean_json_text(text: str) -> str:
    text = text.strip()

    if "```json" in text:
        text = text.split("```json", 1)[1]
    elif "```" in text:
        text = text.split("```", 1)[1]

    if "```" in text:
        text = text.split("```", 1)[0]

    start = text.find("{")
    end = text.rfind("}")

    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]

    return text

def read_pdf(file_path: str):
    reader = PdfReader(file_path)
    resume_data = "".join([page.extract_text() for page in reader.pages])
    return resume_data

def read_docx(file_path: str):
    reader = Document(file_path)
    resume_data = ""
    
    for para in reader.paragraphs:
        resume_data += para.text

    for table in reader.tables:
        for row in table.rows:
            for cell in row.cells:
                resume_data += cell.text
    
    return resume_data

def read_resume(file_path: str): 
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return read_docx(file_path)
    else:
        return None

def getJson(response: str, StructuredClass):

    try:
        response = clean_json_text(response)

        data = json.loads(response)

    except json.JSONDecodeError as e:
        print("\nInvalid JSON returned by LLM")
        print(e)
        print(response)
        return None

    try:
        return StructuredClass.model_validate(data)

    except ValidationError as e:

        print("\nValidation Error")
        print(e)

        print("\nLLM Returned:")

        print(json.dumps(data, indent=4))

        return None